from docx import Document
from typing import List, Dict


def format_authors(authors):
    if not authors:
        return ""

    if len(authors) == 1:
        return authors[0]

    if len(authors) == 2:
        return f"{authors[0]} & {authors[1]}"

    return ", ".join(authors[:-1]) + f", & {authors[-1]}"


def to_apa(paper: Dict) -> str:
    authors = format_authors(paper.get("authors", []))
    year = paper.get("year", "n.d.")
    title = paper.get("title", "")
    journal = paper.get("journal", "")
    doi = paper.get("doi", "")

    text = f"{authors} ({year}). {title}. {journal}."

    if doi:
        doi_clean = doi.replace("https://doi.org/", "")
        text += f" https://doi.org/{doi_clean}"

    return text


def export_apa_docx(manuscript_text: str, papers: List[Dict], output_path: str):
    doc = Document()

    doc.add_heading("Manuscript", level=1)
    doc.add_paragraph(manuscript_text)

    doc.add_heading("References (APA Style)", level=1)

    seen = set()
    merged = []

    for p in papers:
        key = p.get("doi") or p.get("title")
        if key in seen:
            continue
        seen.add(key)
        merged.append(p)

    for i, paper in enumerate(merged, 1):
        doc.add_paragraph(f"{i}. {to_apa(paper)}")

    doc.save(output_path)
    return output_path
